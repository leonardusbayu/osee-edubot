"""
TOEFL ITP Comprehensive Re-import — Correct Part Structure
Format: Options listed without (A)/(B)/(C)/(D) labels, just 4 lines per question
"""
import json, sys, re, os, urllib.request
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
import fitz
try:
    import openpyxl
except:
    pass

DB_ID = 'd501b671-128e-4a45-9d90-74b22e6691ce'
ACCOUNT_ID = '6d55097e8961e881b77a4108e1ae40c2'
API_TOKEN = 'cfut_28l0tEkYzywTicfRbbRcRAl8VLuvTDhBfItQPB5Udca04c87'
BASE = 'D:/TOEFL/MATERIALS'
inserted = errors = 0

def query_d1(sql, params=None):
    url = f'https://api.cloudflare.com/client/v4/accounts/{ACCOUNT_ID}/d1/database/{DB_ID}/query'
    body = {'sql': sql}
    if params: body['params'] = params
    req = urllib.request.Request(url, data=json.dumps(body).encode(),
        headers={'Authorization': f'Bearer {API_TOKEN}', 'Content-Type': 'application/json'})
    try: return json.loads(urllib.request.urlopen(req).read())
    except Exception as e: print(f'    DB: {e}'); return None

def insert_q(section, qtype, title, content):
    global inserted, errors
    r = query_d1('INSERT INTO test_contents (test_type, section, question_type, title, content, media_url, difficulty, topic, source, status) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)',
        ['TOEFL_ITP', section, qtype, title, json.dumps(content, ensure_ascii=False), '', '3', title, 'curated', 'published'])
    if r and r.get('success'): inserted += 1
    else: errors += 1

def read_docx_safe(path):
    if not os.path.exists(path): return []
    try:
        doc = Document(path)
        return [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    except Exception as e:
        print(f'  DOCX error {os.path.basename(path)}: {e}')
        # Try PDF version
        pdf = path.replace('.docx', '.pdf')
        if os.path.exists(pdf):
            try:
                d = fitz.open(pdf)
                text = ''.join([d[i].get_text() for i in range(d.page_count)])
                d.close()
                return [l.strip() for l in text.split('\n') if l.strip()]
            except: pass
        return []

def read_answer_key(path):
    answers = {}
    if not os.path.exists(path): return answers
    try:
        doc = Document(path)
        text = '\n'.join([p.text for p in doc.paragraphs])
        for m in re.finditer(r'(\d+)\.\s*([A-Da-d])', text):
            answers[int(m.group(1))] = m.group(2).upper()
    except: pass
    return answers

def parse_paket(lines, section_name):
    """Parse the Paket format: options listed without (A)(B)(C)(D) labels.
    For listening: 'Options' header → 4 options
    For structure: question line (with ___) → 4 option lines
    For written expression: sentence with underlines → 4 option labels already (A)(B)(C)(D)
    For reading: passage + numbered questions with (A)(B)(C)(D)
    """
    questions = []

    if section_name == 'listening':
        # Format: "Options" → 4 lines of options
        q_num = 0
        i = 0
        while i < len(lines):
            if lines[i].lower().startswith('options') or lines[i].lower() == 'options':
                q_num += 1
                opts = []
                for j in range(1, 5):
                    if i + j < len(lines) and not lines[i+j].lower().startswith('options') and not lines[i+j].lower().startswith('direction'):
                        opts.append(lines[i+j][:150])
                if len(opts) >= 2:
                    questions.append({
                        'num': q_num,
                        'text': '(Listen to the audio)',
                        'options': [{'key': chr(65+k), 'text': o} for k, o in enumerate(opts)]
                    })
                i += len(opts) + 1
            else:
                i += 1

    elif section_name == 'structure':
        # Format: question with ___ → 4 option lines (no A/B/C/D prefix)
        i = 0
        q_num = 0
        while i < len(lines):
            line = lines[i]
            # Skip directions, examples, headers
            if any(x in line.lower() for x in ['direction', 'example', 'sample answer', 'section', 'structure', 'written expression']):
                i += 1
                continue
            # Question: contains ___ or is a sentence that starts a new question
            if '___' in line or '__' in line or ('___ ' in line):
                q_num += 1
                opts = []
                for j in range(1, 5):
                    if i + j < len(lines):
                        opt = lines[i+j]
                        if '___' in opt or '__' in opt or opt.lower().startswith('direction'):
                            break
                        opts.append(opt[:150])
                if len(opts) >= 3:
                    questions.append({
                        'num': q_num,
                        'text': line[:200],
                        'options': [{'key': chr(65+k), 'text': o} for k, o in enumerate(opts)]
                    })
                    i += len(opts) + 1
                else:
                    i += 1
            else:
                i += 1

    elif section_name == 'reading':
        # MCQ format with (A)/(B)/(C)/(D) or just 4 option lines after a question
        current_passage = []
        current_qs = []
        i = 0
        while i < len(lines):
            line = lines[i]
            # Check if it's a question (starts with reasonable text, followed by options)
            # Look for (A) pattern
            if i + 4 < len(lines) and any(lines[i+j].strip().startswith('(') for j in range(1, 5)):
                opts = []
                for j in range(1, 5):
                    if i + j < len(lines):
                        m = re.match(r'^\(?([A-Da-d])\)?\s*(.+)', lines[i+j])
                        if m:
                            opts.append({'key': m.group(1).upper(), 'text': m.group(2)[:150]})
                if len(opts) >= 3:
                    current_qs.append({
                        'num': len(current_qs) + 1,
                        'text': line[:200],
                        'options': opts
                    })
                    i += len(opts) + 1
                    continue
            current_passage.append(line)
            i += 1

        if current_qs:
            questions = current_qs

    return questions

def parse_written_expression(lines):
    """Parse Written Expression (error ID): sentences with underlined parts"""
    questions = []
    i = 0
    q_num = 0
    while i < len(lines):
        line = lines[i]
        if any(x in line.lower() for x in ['direction', 'example', 'sample', 'section', 'written expression']):
            i += 1
            continue
        # In written expression, each question is a sentence — options are the underlined parts
        # The sentence itself contains (A)(B)(C)(D) markers
        if re.search(r'\([A-D]\)', line):
            q_num += 1
            # Extract parts between (A)(B)(C)(D)
            parts = re.split(r'\(([A-D])\)', line)
            opts = []
            full_text = line
            for j in range(1, len(parts), 2):
                if j < len(parts) and j+1 < len(parts):
                    opts.append({'key': parts[j], 'text': parts[j+1].strip()[:100]})
            if len(opts) >= 3:
                questions.append({
                    'num': q_num,
                    'text': full_text[:250],
                    'options': opts
                })
        elif len(line) > 30 and q_num > 0:
            # Might be continuation or new sentence without markers
            pass
        i += 1
    return questions


print('=' * 60)
print('TOEFL ITP COMPREHENSIVE RE-IMPORT')
print('=' * 60)

# ============================================================
# PARSE PAKET 1, 2, 3
# ============================================================
for paket in [1, 2, 3]:
    print(f'\n=== Paket {paket} ===')
    qfile = f'{BASE}/TOEFL_{paket}/Paket {paket}.docx'
    if not os.path.exists(qfile):
        qfile = f'{BASE}/TOEFL_{paket}/paket {paket}.docx'
    afile = f'{BASE}/TOEFL_{paket}/ANSWER KEY PAKET_{paket}.docx'

    lines = read_docx_safe(qfile)
    answers = read_answer_key(afile)
    if not lines:
        print('  No content')
        continue

    # Find section boundaries
    listen_start = listen_end = struct_start = struct_end = we_start = we_end = read_start = 0
    for idx, line in enumerate(lines):
        ll = line.lower()
        if 'section 1' in ll or ('listening' in ll and 'comprehension' in ll):
            listen_start = idx
        if 'section 2' in ll:
            listen_end = idx; struct_start = idx
        if 'written expression' in ll and struct_start > 0:
            we_start = idx
        if 'section 3' in ll or ('reading' in ll and 'comprehension' in ll and idx > struct_start + 5):
            struct_end = idx; read_start = idx

    print(f'  Sections: Listen={listen_start}-{listen_end} Struct={struct_start}-{we_start} WE={we_start}-{struct_end} Read={read_start}-{len(lines)}')

    # Parse Listening
    listen_lines = lines[listen_start:listen_end] if listen_end > listen_start else []
    listen_qs = parse_paket(listen_lines, 'listening')
    print(f'  Listening: {len(listen_qs)} Qs')

    # Part A: 1-30
    for q in [x for x in listen_qs if x['num'] <= 30]:
        title = f'ITP P{paket} L-A Q{q["num"]}'
        content = {'type': 'grouped_listening', 'group_name': f'Paket {paket} Part A Q{q["num"]}', 'part': 'A',
                  'passage_script': '', 'direction': 'Listen and choose the best answer.',
                  'questions': [{'index': 0, 'question_text': q['text'], 'options': q['options'],
                                'answers': [answers.get(q['num'], '')], 'explanation': ''}], 'question_count': 1}
        insert_q('listening', 'listen_short_dialogue', title, content)

    # Part B: 31-38 grouped
    for s, e in [(31, 34), (35, 38)]:
        group = [x for x in listen_qs if s <= x['num'] <= e]
        if not group: continue
        title = f'ITP P{paket} L-B Q{s}-{e}'
        qs = [{'index': i, 'question_text': q['text'], 'options': q['options'],
              'answers': [answers.get(q['num'], '')], 'explanation': ''} for i, q in enumerate(group)]
        content = {'type': 'grouped_listening', 'group_name': f'Paket {paket} Part B', 'part': 'B',
                  'passage_script': '', 'direction': 'Listen to the conversation.',
                  'questions': qs, 'question_count': len(qs)}
        insert_q('listening', 'listen_long_conversation', title, content)

    # Part C: 39-50 grouped
    c_groups = [(38, 41), (42, 45), (46, 50)] if paket == 3 else [(39, 42), (43, 46), (47, 50)]
    for s, e in c_groups:
        group = [x for x in listen_qs if s <= x['num'] <= e]
        if not group: continue
        title = f'ITP P{paket} L-C Q{s}-{e}'
        qs = [{'index': i, 'question_text': q['text'], 'options': q['options'],
              'answers': [answers.get(q['num'], '')], 'explanation': ''} for i, q in enumerate(group)]
        content = {'type': 'grouped_listening', 'group_name': f'Paket {paket} Part C', 'part': 'C',
                  'passage_script': '', 'direction': 'Listen to the talk.',
                  'questions': qs, 'question_count': len(qs)}
        insert_q('listening', 'listen_talk', title, content)

    # Parse Structure Part A
    struct_lines = lines[struct_start:we_start] if we_start > struct_start else lines[struct_start:struct_end]
    struct_qs = parse_paket(struct_lines, 'structure')
    print(f'  Structure A: {len(struct_qs)} Qs')

    for i, q in enumerate(struct_qs):
        title = f'ITP P{paket} S-A Q{i+1}'
        content = {'type': 'grouped_reading', 'group_name': f'Paket {paket} Structure Part A', 'passage': '',
                  'direction': 'Complete the sentence.',
                  'questions': [{'index': 0, 'question_text': q['text'], 'options': q['options'],
                                'answers': [answers.get(q['num'] + 50 if paket <= 3 else q['num'], '')], 'explanation': ''}],
                  'question_count': 1}
        insert_q('structure', 'sentence_completion', title, content)

    # Parse Written Expression Part B
    we_lines = lines[we_start:struct_end] if struct_end > we_start else []
    we_qs = parse_written_expression(we_lines)
    if not we_qs:
        # Fallback: parse as structure-style
        we_qs = parse_paket(we_lines, 'structure')
    print(f'  Written Expression B: {len(we_qs)} Qs')

    for i, q in enumerate(we_qs):
        title = f'ITP P{paket} S-B Q{i+16}'
        content = {'type': 'grouped_reading', 'group_name': f'Paket {paket} Written Expression Part B', 'passage': '',
                  'direction': 'Identify the underlined word or phrase that must be changed.',
                  'questions': [{'index': 0, 'question_text': q['text'], 'options': q['options'],
                                'answers': [answers.get(q['num'] + 65 if paket <= 3 else q['num'], '')], 'explanation': ''}],
                  'question_count': 1}
        insert_q('structure', 'error_identification', title, content)

    # Parse Reading
    read_lines = lines[read_start:] if read_start > 0 else []
    read_qs = parse_paket(read_lines, 'reading')
    print(f'  Reading: {len(read_qs)} Qs')

    for i in range(0, len(read_qs), 8):
        batch = read_qs[i:i+8]
        title = f'ITP P{paket} R {i//8+1}'
        qs = [{'index': j, 'question_text': q['text'], 'options': q['options'],
              'answers': [answers.get(q['num'] + 90 if paket <= 3 else q['num'], '')], 'explanation': ''} for j, q in enumerate(batch)]
        content = {'type': 'grouped_reading', 'group_name': f'Paket {paket} Reading', 'passage': '',
                  'direction': 'Read the passage and answer.',
                  'questions': qs, 'question_count': len(qs)}
        insert_q('reading', 'reading_passage', title, content)


# ============================================================
# PARSE PRE-TEST + POST-TEST (separate DOCX per section)
# ============================================================
for test_name, prefix, section_files in [
    ('Pre-test', 'PRE', {
        'structure': f'{BASE}/TOEFL Pre-test/STRUCTURE AND WRITTEN DIAGNOSTIC TEST',
        'reading': f'{BASE}/TOEFL Pre-test/READING DIAGNOSTIC TEST',
    }),
    ('Post-test', 'POST', {
        'structure': f'{BASE}/TOEFL Post-test/STRUCTURE AND WRITTEN POST TEST',
        'reading': f'{BASE}/TOEFL Post-test/READING POST TEST',
    }),
]:
    print(f'\n=== {test_name} ===')
    for section, base_path in section_files.items():
        # Try DOCX first, then PDF
        lines = read_docx_safe(base_path + '.docx')
        if not lines:
            lines = read_docx_safe(base_path + '.pdf')  # Uses PDF fallback
        if not lines:
            print(f'  {section}: no content')
            continue

        if section == 'structure':
            # Find Written Expression boundary
            we_idx = 0
            for idx, line in enumerate(lines):
                if 'written expression' in line.lower():
                    we_idx = idx
                    break

            # Part A
            sa_lines = lines[:we_idx] if we_idx > 0 else lines[:len(lines)//2]
            sa_qs = parse_paket(sa_lines, 'structure')
            print(f'  Structure A: {len(sa_qs)} Qs')
            for i, q in enumerate(sa_qs):
                title = f'ITP {prefix} S-A Q{i+1}'
                content = {'type': 'grouped_reading', 'group_name': f'{test_name} Structure A', 'passage': '',
                          'direction': 'Complete the sentence.',
                          'questions': [{'index': 0, 'question_text': q['text'], 'options': q['options'], 'answers': [], 'explanation': ''}],
                          'question_count': 1}
                insert_q('structure', 'sentence_completion', title, content)

            # Part B
            sb_lines = lines[we_idx:] if we_idx > 0 else lines[len(lines)//2:]
            sb_qs = parse_written_expression(sb_lines)
            if not sb_qs:
                sb_qs = parse_paket(sb_lines, 'structure')
            print(f'  Written Expression B: {len(sb_qs)} Qs')
            for i, q in enumerate(sb_qs):
                title = f'ITP {prefix} S-B Q{i+16}'
                content = {'type': 'grouped_reading', 'group_name': f'{test_name} Written Expression B', 'passage': '',
                          'direction': 'Find the error.',
                          'questions': [{'index': 0, 'question_text': q['text'], 'options': q['options'], 'answers': [], 'explanation': ''}],
                          'question_count': 1}
                insert_q('structure', 'error_identification', title, content)

        elif section == 'reading':
            read_qs = parse_paket(lines, 'reading')
            print(f'  Reading: {len(read_qs)} Qs')
            for i in range(0, len(read_qs), 8):
                batch = read_qs[i:i+8]
                title = f'ITP {prefix} R {i//8+1}'
                qs = [{'index': j, 'question_text': q['text'], 'options': q['options'], 'answers': [], 'explanation': ''} for j, q in enumerate(batch)]
                content = {'type': 'grouped_reading', 'group_name': f'{test_name} Reading', 'passage': '',
                          'direction': 'Read and answer.', 'questions': qs, 'question_count': len(qs)}
                insert_q('reading', 'reading_passage', title, content)


# ============================================================
# PARSE COMPILATION
# ============================================================
print('\n=== Compilation ===')
for filepath, section, qtype in [
    (f'{BASE}/TOEFL Compilation/Structure/toefl structure.docx', 'structure', 'sentence_completion'),
    (f'{BASE}/TOEFL Compilation/TOEFL Handout plus answer key.docx', 'structure', 'sentence_completion'),
]:
    lines = read_docx_safe(filepath)
    if not lines: continue
    qs = parse_paket(lines, 'structure')
    fname = os.path.basename(filepath)[:25]
    print(f'  {fname}: {len(qs)} Qs')
    for i in range(0, len(qs), 10):
        batch = qs[i:i+10]
        title = f'ITP COMP {fname} {i//10+1}'
        gqs = [{'index': j, 'question_text': q['text'], 'options': q['options'], 'answers': [], 'explanation': ''} for j, q in enumerate(batch)]
        content = {'type': 'grouped_reading', 'group_name': title, 'passage': '', 'direction': 'Choose the best answer.',
                  'questions': gqs, 'question_count': len(gqs)}
        insert_q(section, qtype, title, content)


# ============================================================
# miniTOEFL from PDF
# ============================================================
print('\n=== miniTOEFL ===')
mini_pdf = f'{BASE}/miniTOEFL v.1.0/Mini TOEFL v.1.0.pdf'
if os.path.exists(mini_pdf):
    d = fitz.open(mini_pdf)
    text = ''.join([d[i].get_text() for i in range(d.page_count)])
    d.close()
    lines = [l.strip() for l in text.split('\n') if l.strip()]

    # Structure section
    struct_qs = parse_paket(lines, 'structure')
    print(f'  Structure from PDF: {len(struct_qs)} Qs')
    for i, q in enumerate(struct_qs):
        qtype = 'sentence_completion' if i < 15 else 'error_identification'
        title = f'ITP MINI S Q{i+1}'
        content = {'type': 'grouped_reading', 'group_name': f'miniTOEFL Structure', 'passage': '',
                  'direction': 'Choose the best answer.',
                  'questions': [{'index': 0, 'question_text': q['text'], 'options': q['options'], 'answers': [], 'explanation': ''}],
                  'question_count': 1}
        insert_q('structure', qtype, title, content)


# ============================================================
print(f'\n{"="*60}')
print(f'TOTAL: {inserted} groups inserted | {errors} errors')
print(f'{"="*60}')
