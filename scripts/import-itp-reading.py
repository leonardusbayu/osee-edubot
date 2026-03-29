"""
Import TOEFL ITP Reading Comprehension from Paket 1/2/3 + Pre-test + Post-test DOCXs.
Extracts passage text + questions grouped by "Questions X-Y" headers.
Also applies answer keys.
"""
import json, sys, re, urllib.request
sys.stdout.reconfigure(encoding='utf-8')

try:
    import docx
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx', 'openpyxl'])
    import docx

import openpyxl

DB_ID = 'd501b671-128e-4a45-9d90-74b22e6691ce'
ACCOUNT_ID = '6d55097e8961e881b77a4108e1ae40c2'
API_TOKEN = 'cfut_28l0tEkYzywTicfRbbRcRAl8VLuvTDhBfItQPB5Udca04c87'
BASE = 'D:/TOEFL/MATERIALS'

def query_d1(sql, params=None):
    url = f'https://api.cloudflare.com/client/v4/accounts/{ACCOUNT_ID}/d1/database/{DB_ID}/query'
    body = {'sql': sql}
    if params:
        body['params'] = params
    req = urllib.request.Request(url, data=json.dumps(body).encode(),
        headers={'Authorization': f'Bearer {API_TOKEN}', 'Content-Type': 'application/json'})
    try:
        resp = urllib.request.urlopen(req)
        result = json.loads(resp.read())
        return result
    except Exception as e:
        print(f'  D1 error: {e}')
        return None

def insert_q(section, qtype, title, content):
    r = query_d1(
        'INSERT INTO test_contents (test_type, section, question_type, title, content, media_url, difficulty, topic, source, status) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)',
        ['TOEFL_ITP', section, qtype, title, json.dumps(content, ensure_ascii=False), '', '3', title, 'curated', 'published']
    )
    ok = r and r.get('success')
    if not ok:
        print(f'  INSERT FAILED for {title}')
    return ok

def delete_existing():
    """Delete existing ITP reading_passage entries to avoid duplicates."""
    r = query_d1("DELETE FROM test_contents WHERE test_type = 'TOEFL_ITP' AND question_type = 'reading_passage'")
    if r and r.get('success'):
        info = r.get('result', [{}])
        if isinstance(info, list) and len(info) > 0:
            changes = info[0].get('meta', {}).get('changes', 0)
            print(f'Deleted {changes} existing reading_passage rows')
    else:
        print('Warning: could not delete existing rows')


# ============================================================
# PARSE ANSWER KEYS
# ============================================================

def parse_answer_key_table(doc_path):
    """Parse answer key from DOCX tables (Paket 1/2/3 format).
    Table 2 (index 2) = Reading answers, 50 questions across 5 columns of 10."""
    d = docx.Document(doc_path)
    answers = {}
    if len(d.tables) < 3:
        print(f'  Warning: {doc_path} has only {len(d.tables)} tables, expected 3')
        return answers
    table = d.tables[2]  # Reading table
    for row in table.rows:
        for cell in row.cells:
            text = cell.text.strip()
            for line in text.split('\n'):
                line = line.strip()
                m = re.match(r'(\d+)\.([A-D])', line)
                if m:
                    answers[int(m.group(1))] = m.group(2)
    return answers

def parse_answer_key_xlsx(sheet_name, col_index):
    """Parse answer key from Excel file. col_index is 0-based column for reading answers."""
    wb = openpyxl.load_workbook(f'{BASE}/Answer Key For TOEFL_Revised.xlsx')
    ws = wb[sheet_name]
    answers = {}
    for i, row in enumerate(ws.iter_rows(min_row=1, max_row=50, values_only=True)):
        if col_index < len(row) and row[col_index]:
            answers[i + 1] = str(row[col_index]).strip()
    return answers


# ============================================================
# PARSE READING FROM PAKET DOCX (Questions X-Y format)
# ============================================================

def parse_paket_reading(doc_path, answers, paket_num):
    """Parse reading section from Paket DOCX.
    Key insight: passage paragraphs use 'Normal' style, while questions and options
    use 'List Paragraph' style. We split on that boundary.
    Within the List Paragraph block, every 5 consecutive lines = 1 question + 4 options.
    """
    d = docx.Document(doc_path)
    paragraphs = [(p.text.strip(), p.style.name) for p in d.paragraphs]

    # Find reading section start
    reading_start = 0
    for i, (text, style) in enumerate(paragraphs):
        if 'reading comprehension' in text.lower() and len(text) < 50:
            reading_start = i
            break

    # Find all "Questions X-Y" headers after reading_start
    question_groups = []
    for i, (text, style) in enumerate(paragraphs):
        if i < reading_start:
            continue
        m = re.match(r'Questions?\s+(\d+)\s*[-–]\s*(\d+)', text, re.IGNORECASE)
        if m:
            question_groups.append({
                'start_idx': i,
                'q_start': int(m.group(1)),
                'q_end': int(m.group(2)),
                'header': text,
            })

    if not question_groups:
        print(f'  No question groups found in {doc_path}')
        return []

    results = []
    for g_idx, group in enumerate(question_groups):
        end_idx = question_groups[g_idx + 1]['start_idx'] if g_idx + 1 < len(question_groups) else len(paragraphs)

        # Collect all non-empty paragraphs with styles
        all_paras = []
        for i in range(group['start_idx'] + 1, end_idx):
            text, style = paragraphs[i]
            if text:
                all_paras.append((text, style))

        expected_count = group['q_end'] - group['q_start'] + 1

        # Strategy: Use style-based split. Passage = Normal paragraphs before
        # the question block. Question block = consecutive List Paragraph lines.
        # Edge case: first passage paragraph may also be List Paragraph (e.g. P1 Q12-23).
        # Solution: find the LAST Normal paragraph, then look for the first
        # List Paragraph AFTER that which starts the question block.

        # Find where Normal paragraphs end (passage body)
        last_normal_idx = -1
        for idx, (text, style) in enumerate(all_paras):
            if 'list' not in style.lower() and not style.startswith('Style'):
                last_normal_idx = idx

        # Questions start at first List Paragraph after last Normal paragraph
        if last_normal_idx >= 0:
            boundary = last_normal_idx + 1
        else:
            # All List Paragraph — use expected_count * 5 from end
            boundary = max(0, len(all_paras) - expected_count * 5)

        passage_lines = [t for t, s in all_paras[:boundary]]
        list_lines = [t for t, s in all_paras[boundary:]]

        # Clean passage
        passage_text = '\n'.join(passage_lines).strip()
        passage_text = re.sub(r'^(As soon as you understand.*?section\.?\s*)', '', passage_text, flags=re.IGNORECASE)

        # Parse questions from list_lines: every 5 lines = question + 4 options
        expected_count = group['q_end'] - group['q_start'] + 1
        questions = []

        # Clean list lines: some options have (A)/(B)/(C)/(D) prefix
        i = 0
        while i < len(list_lines):
            # Look for question text (not an option label)
            q_text = list_lines[i].strip()
            # Remove leading number if present
            q_text = re.sub(r'^\d+\.\s*', '', q_text).strip()

            if i + 4 < len(list_lines):
                opts_raw = list_lines[i+1:i+5]
            elif i + 4 == len(list_lines):
                opts_raw = list_lines[i+1:i+5]
            else:
                break

            # Clean options
            clean_opts = []
            for idx, o in enumerate(opts_raw):
                letter = chr(65 + idx)
                cleaned = re.sub(r'^\(?[A-D]\)?\s*', '', o).strip()
                cleaned = re.sub(r'^\(?[A-D]\)?\t\s*', '', cleaned).strip()
                clean_opts.append({'key': letter, 'text': cleaned})

            q_num = group['q_start'] + len(questions)
            ans = answers.get(q_num, '')
            questions.append({
                'index': len(questions),
                'question_text': q_text,
                'options': clean_opts,
                'answers': [ans] if ans else [],
                'explanation': '',
            })

            i += 5
            if len(questions) >= expected_count:
                break

        if not questions:
            print(f'  Warning: no questions extracted for {group["header"]}')
            continue

        title = f'ITP P{paket_num} Reading Q{group["q_start"]}-{group["q_end"]}'
        results.append({
            'title': title,
            'group_name': f'Paket {paket_num} Reading: Questions {group["q_start"]}-{group["q_end"]}',
            'passage': passage_text,
            'questions': questions,
            'q_start': group['q_start'],
            'q_end': group['q_end'],
        })

    return results


# ============================================================
# PARSE READING FROM PRE-TEST/POST-TEST DOCX (numbered questions format)
# ============================================================

def parse_prepost_reading(doc_path, answers, test_name, prefix):
    """Parse reading from Pre-test/Post-test format.
    Format: 'Questions X-Y' header -> passage -> numbered questions (1. ...) with 4 options.
    Options may have (A)/(B)/(C)/(D) labels or just be plain lines in List Paragraph style.
    """
    d = docx.Document(doc_path)
    paragraphs = [(p.text.strip(), p.style.name) for p in d.paragraphs]

    # Find all "Questions X-Y" headers
    question_groups = []
    for i, (text, style) in enumerate(paragraphs):
        m = re.match(r'Questions?\s+(\d+)\s*[-–]\s*(\d+)', text, re.IGNORECASE)
        if m:
            question_groups.append({
                'start_idx': i,
                'q_start': int(m.group(1)),
                'q_end': int(m.group(2)),
                'header': text,
            })

    if not question_groups:
        print(f'  No question groups found in {doc_path}')
        return []

    results = []
    for g_idx, group in enumerate(question_groups):
        end_idx = question_groups[g_idx + 1]['start_idx'] if g_idx + 1 < len(question_groups) else len(paragraphs)

        # Collect paragraphs
        group_paras = []
        for i in range(group['start_idx'] + 1, end_idx):
            text, style = paragraphs[i]
            if text:
                group_paras.append((text, style))

        # In pre/post format, questions start with "N. " (number dot space)
        # Find first numbered question
        first_q_idx = None
        for i, (text, style) in enumerate(group_paras):
            if re.match(r'^\d+\.\s', text):
                first_q_idx = i
                break

        if first_q_idx is None:
            # Try alternative: look for question-like text
            for i, (text, style) in enumerate(group_paras):
                if text.endswith('?') or 'closest in meaning' in text.lower():
                    first_q_idx = i
                    break

        if first_q_idx is None:
            print(f'  Warning: no questions found for {group["header"]} in {doc_path}')
            continue

        # Passage = everything before first question
        passage_text = '\n'.join(t for t, s in group_paras[:first_q_idx]).strip()

        # Parse questions
        q_paras = group_paras[first_q_idx:]
        questions = []
        current_q = None
        current_opts = []

        for text, style in q_paras:
            # Check if it's a new numbered question
            m = re.match(r'^(\d+)\.\s+(.*)', text)
            if m:
                # Save previous question
                if current_q is not None and len(current_opts) >= 4:
                    q_num = group['q_start'] + len(questions)
                    ans = answers.get(q_num, '')
                    questions.append({
                        'index': len(questions),
                        'question_text': current_q,
                        'options': current_opts[:4],
                        'answers': [ans] if ans else [],
                        'explanation': '',
                    })
                current_q = m.group(2).strip()
                current_opts = []
            elif current_q is not None:
                # This is an option line
                # Remove (A)/(B)/(C)/(D) prefix if present
                opt_text = re.sub(r'^\(?[A-D]\)?\s*', '', text).strip()
                if opt_text:
                    letter = chr(65 + len(current_opts))
                    current_opts.append({'key': letter, 'text': opt_text})

        # Save last question
        if current_q is not None and len(current_opts) >= 4:
            q_num = group['q_start'] + len(questions)
            ans = answers.get(q_num, '')
            questions.append({
                'index': len(questions),
                'question_text': current_q,
                'options': current_opts[:4],
                'answers': [ans] if ans else [],
                'explanation': '',
            })

        if not questions:
            print(f'  Warning: no questions parsed for {group["header"]}')
            continue

        # Fix indices
        for i, q in enumerate(questions):
            q['index'] = i

        title = f'ITP {prefix} Reading Q{group["q_start"]}-{group["q_end"]}'
        results.append({
            'title': title,
            'group_name': f'{test_name} Reading: Questions {group["q_start"]}-{group["q_end"]}',
            'passage': passage_text,
            'questions': questions,
            'q_start': group['q_start'],
            'q_end': group['q_end'],
        })

    return results


# ============================================================
# MAIN
# ============================================================

inserted = 0

# Delete existing reading_passage entries
delete_existing()

# --- PAKET 1, 2, 3 ---
paket_configs = [
    (1, f'{BASE}/TOEFL_1/Paket 1.docx', f'{BASE}/TOEFL_1/ANSWER KEY PAKET_1.docx'),
    (2, f'{BASE}/TOEFL_2/Paket 2.docx', f'{BASE}/TOEFL_2/ANSWER KEY PAKET_2.docx'),
    (3, f'{BASE}/TOEFL_3/paket 3.docx', f'{BASE}/TOEFL_3/ANSWER KEY PAKET_3.docx'),
]

for paket_num, doc_path, key_path in paket_configs:
    print(f'\n=== Paket {paket_num} ===')
    answers = parse_answer_key_table(key_path)
    print(f'  Answer key: {len(answers)} answers')

    groups = parse_paket_reading(doc_path, answers, paket_num)
    for g in groups:
        content = {
            'type': 'grouped_reading',
            'group_name': g['group_name'],
            'passage': g['passage'],
            'direction': 'Read the passage and choose the best answer.',
            'questions': g['questions'],
            'question_count': len(g['questions']),
        }
        if insert_q('reading', 'reading_passage', g['title'], content):
            inserted += 1
            q_with_ans = sum(1 for q in g['questions'] if q['answers'])
            print(f"  {g['title']}: {len(g['questions'])} Qs ({q_with_ans} with answers), passage {len(g['passage'])} chars")


# --- PRE-TEST ---
print('\n=== Pre-test Reading ===')
pretest_answers = parse_answer_key_xlsx('TOEFL PRE TEST', 2)  # Column 3 (0-indexed: 2)
print(f'  Answer key: {len(pretest_answers)} answers')

pretest_groups = parse_prepost_reading(
    f'{BASE}/TOEFL Pre-test/READING DIAGNOSTIC TEST.docx',
    pretest_answers, 'Pre-test', 'PRE'
)
for g in pretest_groups:
    content = {
        'type': 'grouped_reading',
        'group_name': g['group_name'],
        'passage': g['passage'],
        'direction': 'Read the passage and choose the best answer.',
        'questions': g['questions'],
        'question_count': len(g['questions']),
    }
    if insert_q('reading', 'reading_passage', g['title'], content):
        inserted += 1
        q_with_ans = sum(1 for q in g['questions'] if q['answers'])
        print(f"  {g['title']}: {len(g['questions'])} Qs ({q_with_ans} with answers), passage {len(g['passage'])} chars")


# --- POST-TEST ---
print('\n=== Post-test Reading ===')
posttest_answers = parse_answer_key_xlsx('TOEFL POST TEST', 2)  # Column 3 (0-indexed: 2)
print(f'  Answer key: {len(posttest_answers)} answers')

posttest_groups = parse_prepost_reading(
    f'{BASE}/TOEFL Post-test/READING POST TEST.docx',
    posttest_answers, 'Post-test', 'POST'
)
for g in posttest_groups:
    content = {
        'type': 'grouped_reading',
        'group_name': g['group_name'],
        'passage': g['passage'],
        'direction': 'Read the passage and choose the best answer.',
        'questions': g['questions'],
        'question_count': len(g['questions']),
    }
    if insert_q('reading', 'reading_passage', g['title'], content):
        inserted += 1
        q_with_ans = sum(1 for q in g['questions'] if q['answers'])
        print(f"  {g['title']}: {len(g['questions'])} Qs ({q_with_ans} with answers), passage {len(g['passage'])} chars")


print(f'\n=== TOTAL: {inserted} reading groups inserted ===')
