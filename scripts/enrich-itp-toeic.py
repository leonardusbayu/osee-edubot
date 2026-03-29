"""Enrich TOEFL ITP + TOEIC from all available materials"""
import json, sys, re, os, urllib.request
sys.stdout.reconfigure(encoding='utf-8')
import fitz
from docx import Document

DB_ID = 'd501b671-128e-4a45-9d90-74b22e6691ce'
ACCOUNT_ID = '6d55097e8961e881b77a4108e1ae40c2'
API_TOKEN = 'cfut_28l0tEkYzywTicfRbbRcRAl8VLuvTDhBfItQPB5Udca04c87'

def query_d1(sql, params=None):
    url = f'https://api.cloudflare.com/client/v4/accounts/{ACCOUNT_ID}/d1/database/{DB_ID}/query'
    body = {'sql': sql}
    if params: body['params'] = params
    req = urllib.request.Request(url, data=json.dumps(body).encode(), headers={'Authorization': f'Bearer {API_TOKEN}', 'Content-Type': 'application/json'})
    try: resp = urllib.request.urlopen(req); return json.loads(resp.read())
    except: return None

def insert_q(test_type, section, qtype, title, content):
    r = query_d1('INSERT INTO test_contents (test_type, section, question_type, title, content, media_url, difficulty, topic, source, status) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)',
        [test_type, section, qtype, title, json.dumps(content, ensure_ascii=False), '', '3', title, 'curated', 'published'])
    return r and r.get('success')

inserted = 0

# ============================================================
# TOEFL ITP — Parse DOCX test packets
# ============================================================
print('=== TOEFL ITP DOCX Packets ===')

for paket_num in [1, 2, 3]:
    paket_file = f'D:/TOEFL/MATERIALS/TOEFL_{paket_num}/Paket {paket_num}.docx'
    answer_file = f'D:/TOEFL/MATERIALS/TOEFL_{paket_num}/ANSWER KEY PAKET_{paket_num}.docx'

    if not os.path.exists(paket_file):
        paket_file = f'D:/TOEFL/MATERIALS/TOEFL_{paket_num}/paket {paket_num}.docx'
    if not os.path.exists(paket_file):
        print(f'  Paket {paket_num}: file not found')
        continue

    try:
        doc = Document(paket_file)
        text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])

        # Extract MCQ questions: number. question text  (A) opt (B) opt (C) opt (D) opt
        q_pattern = re.compile(r'(\d+)\.\s*(.+?)(?:\n|\s{2,})\(A\)\s*(.+?)(?:\n|\s{2,})\(B\)\s*(.+?)(?:\n|\s{2,})\(C\)\s*(.+?)(?:\n|\s{2,})\(D\)\s*(.+?)(?=\n\d+\.|\n\n|$)', re.DOTALL)
        matches = q_pattern.findall(text)

        if not matches:
            # Try simpler pattern
            lines = text.split('\n')
            current_q = None
            questions = []
            for line in lines:
                line = line.strip()
                m = re.match(r'^(\d+)\.\s+(.+)', line)
                if m:
                    if current_q and len(current_q.get('options', [])) >= 4:
                        questions.append(current_q)
                    current_q = {'num': int(m.group(1)), 'text': m.group(2), 'options': []}
                elif current_q:
                    om = re.match(r'^\(([A-D])\)\s*(.+)', line)
                    if om:
                        current_q['options'].append({'key': om.group(1), 'text': om.group(2)})
            if current_q and len(current_q.get('options', [])) >= 4:
                questions.append(current_q)
        else:
            questions = [{'num': int(n), 'text': q.strip(), 'options': [
                {'key': 'A', 'text': a.strip()}, {'key': 'B', 'text': b.strip()},
                {'key': 'C', 'text': c.strip()}, {'key': 'D', 'text': d.strip()}
            ]} for n, q, a, b, c, d in matches]

        # Read answer key
        answers = {}
        if os.path.exists(answer_file):
            try:
                ak_doc = Document(answer_file)
                ak_text = '\n'.join([p.text for p in ak_doc.paragraphs])
                for m in re.finditer(r'(\d+)\.\s*([A-D])', ak_text):
                    answers[int(m.group(1))] = m.group(2)
            except:
                pass

        # Split into sections: 1-50 Listening, 51-90 Structure, 91-140 Reading
        for section_name, q_range, section_id, qtype in [
            ('Listening', (1, 50), 'listening', 'short_dialogues'),
            ('Structure', (51, 90), 'structure', 'sentence_completion'),
            ('Reading', (91, 140), 'reading', 'reading_comprehension'),
        ]:
            section_qs = [q for q in questions if q_range[0] <= q['num'] <= q_range[1]]
            if not section_qs:
                continue

            # Insert in batches of 10
            for i in range(0, len(section_qs), 10):
                batch = section_qs[i:i+10]
                qs = [{
                    'index': j,
                    'question_text': q['text'][:200],
                    'options': q.get('options', []),
                    'answers': [answers.get(q['num'], '')] if q['num'] in answers else [],
                    'explanation': '',
                } for j, q in enumerate(batch)]

                title = f'ITP Paket {paket_num} {section_name} {i//10+1}'
                content = {
                    'type': 'grouped_reading' if section_id != 'listening' else 'grouped_listening',
                    'group_name': title, 'passage': '', 'direction': 'Choose the best answer.',
                    'questions': qs, 'question_count': len(qs),
                }
                if section_id == 'listening':
                    content['passage_script'] = ''

                if insert_q('TOEFL_ITP', section_id, qtype, title, content):
                    inserted += 1
                    print(f'  {title}: {len(qs)} Qs')

    except Exception as e:
        print(f'  Paket {paket_num} error: {e}')

# ============================================================
# TOEFL ITP — Parse diagnostic/compilation DOCX
# ============================================================
print('\n=== ITP Diagnostic/Compilation ===')

docx_files = [
    ('D:/TOEFL/MATERIALS/TOEFL Compilation/Structure/toefl structure.docx', 'structure', 'sentence_completion'),
    ('D:/TOEFL/MATERIALS/TOEFL Compilation/Reading/READING DIAGNOSTIC TEST.docx', 'reading', 'reading_comprehension'),
    ('D:/TOEFL/MATERIALS/TOEFL Compilation/Structure/STRUCTURE AND WRITTEN DIAGNOSTIC TEST.docx', 'structure', 'written_expression'),
    ('D:/TOEFL/MATERIALS/TOEFL Compilation/TOEFL Handout plus answer key.docx', 'structure', 'sentence_completion'),
    ('D:/TOEFL/MATERIALS/TOEFL Pre-test/STRUCTURE AND WRITTEN DIAGNOSTIC TEST.docx', 'structure', 'sentence_completion'),
    ('D:/TOEFL/MATERIALS/TOEFL Pre-test/READING DIAGNOSTIC TEST.docx', 'reading', 'reading_comprehension'),
    ('D:/TOEFL/MATERIALS/TOEFL Post-test/STRUCTURE AND WRITTEN POST TEST.docx', 'structure', 'sentence_completion'),
    ('D:/TOEFL/MATERIALS/TOEFL Post-test/READING POST TEST.docx', 'reading', 'reading_comprehension'),
]

for filepath, section, qtype in docx_files:
    if not os.path.exists(filepath):
        continue
    try:
        doc = Document(filepath)
        text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])

        # Extract questions
        questions = []
        lines = text.split('\n')
        current_q = None
        for line in lines:
            line = line.strip()
            m = re.match(r'^(\d+)\.\s+(.+)', line)
            if m and len(m.group(2)) > 5:
                if current_q and current_q.get('options'):
                    questions.append(current_q)
                current_q = {'num': int(m.group(1)), 'text': m.group(2)[:200], 'options': []}
            elif current_q:
                om = re.match(r'^\(?([A-Da-d])\)?\s*(.+)', line)
                if om and len(om.group(2)) > 1:
                    current_q['options'].append({'key': om.group(1).upper(), 'text': om.group(2)[:100]})
        if current_q and current_q.get('options'):
            questions.append(current_q)

        fname = os.path.basename(filepath).replace('.docx', '')

        for i in range(0, len(questions), 10):
            batch = questions[i:i+10]
            qs = [{'index': j, 'question_text': q['text'], 'options': q.get('options', []), 'answers': [], 'explanation': ''} for j, q in enumerate(batch)]
            title = f'ITP {fname} {i//10+1}'
            ctype = 'grouped_listening' if section == 'listening' else 'grouped_reading'
            content = {'type': ctype, 'group_name': title, 'passage': '', 'direction': 'Choose the best answer.', 'questions': qs, 'question_count': len(qs)}
            if section == 'listening':
                content['passage_script'] = ''
            if insert_q('TOEFL_ITP', section, qtype, title, content):
                inserted += 1
                print(f'  {title}: {len(qs)} Qs')

    except Exception as e:
        print(f'  {os.path.basename(filepath)}: {e}')

# ============================================================
# TOEFL ITP — Parse PDFs
# ============================================================
print('\n=== ITP PDFs ===')

pdf_files = [
    ('D:/TOEFL/MATERIALS/miniTOEFL v.1.0/Mini TOEFL v.1.0.pdf', 'structure'),
    ('D:/TOEFL/MATERIALS/TOEFL-like Test Sample/TOEFL-like Test_post test longman.pdf', 'structure'),
    ('D:/TOEFL/MATERIALS/TOEFL-like Test Sample/answer key.pdf', None),  # just answers
]

for filepath, default_section in pdf_files:
    if not os.path.exists(filepath) or default_section is None:
        continue
    try:
        doc = fitz.open(filepath)
        text = ''
        for i in range(doc.page_count):
            text += doc[i].get_text() + '\n'
        doc.close()

        if len(text) < 100:
            continue

        # Extract MCQ
        mcq = re.findall(r'(\d+)\.\s+(.+?)\n\s*\(A\)\s*(.+?)\n\s*\(B\)\s*(.+?)\n\s*\(C\)\s*(.+?)\n\s*\(D\)\s*(.+?)(?=\n\d+\.|\n\n|$)', text, re.DOTALL)
        if not mcq:
            # Simpler pattern
            mcq = re.findall(r'(\d+)\.\s+(.{10,200}?)(?:\n|$)', text)
            questions = [{'index': i, 'question_text': q.strip(), 'options': [], 'answers': [], 'explanation': ''} for i, (_, q) in enumerate(mcq[:30]) if len(q.strip()) > 10]
        else:
            questions = [{'index': i, 'question_text': q.strip().replace('\n',' ')[:200], 'options': [{'key':'A','text':a.strip()[:100]},{'key':'B','text':b.strip()[:100]},{'key':'C','text':c.strip()[:100]},{'key':'D','text':d.strip()[:100]}], 'answers': [], 'explanation': ''} for i, (_, q, a, b, c, d) in enumerate(mcq)]

        fname = os.path.basename(filepath).replace('.pdf', '')
        for i in range(0, len(questions), 10):
            batch = questions[i:i+10]
            title = f'ITP {fname} {i//10+1}'
            content = {'type': 'grouped_reading', 'group_name': title, 'passage': '', 'direction': 'Choose the best answer.', 'questions': batch, 'question_count': len(batch)}
            if insert_q('TOEFL_ITP', default_section, 'sentence_completion', title, content):
                inserted += 1
                print(f'  {title}: {len(batch)} Qs')

    except Exception as e:
        print(f'  {os.path.basename(filepath)}: {e}')

# ============================================================
# TOEIC — AI-generate more business English MCQ
# ============================================================
print('\n=== TOEIC AI Generation ===')

# Generate Part 5 style questions programmatically (no API call needed — template-based)
toeic_templates = [
    {'q': 'The annual meeting will be ___ in the main conference hall.', 'A': 'hold', 'B': 'held', 'C': 'holding', 'D': 'holds', 'a': 'B'},
    {'q': 'Applicants must have at least five years of ___ experience.', 'A': 'relate', 'B': 'relative', 'C': 'related', 'D': 'relation', 'a': 'C'},
    {'q': 'The ___ of the building has been delayed due to bad weather.', 'A': 'construct', 'B': 'construction', 'C': 'constructive', 'D': 'constructing', 'a': 'B'},
    {'q': 'Please ensure that all documents are ___ before the deadline.', 'A': 'submit', 'B': 'submitted', 'C': 'submitting', 'D': 'submission', 'a': 'B'},
    {'q': 'The new product line has been ___ well received by customers.', 'A': 'extreme', 'B': 'extremely', 'C': 'extremity', 'D': 'extremes', 'a': 'B'},
    {'q': '___ the increase in demand, the company hired additional staff.', 'A': 'Because', 'B': 'Although', 'C': 'Due to', 'D': 'Despite', 'a': 'C'},
    {'q': 'The warranty ___ defects in materials and workmanship.', 'A': 'cover', 'B': 'covers', 'C': 'covering', 'D': 'covered', 'a': 'B'},
    {'q': 'Employees who work overtime are ___ for additional compensation.', 'A': 'eligible', 'B': 'eligibility', 'C': 'eligibly', 'D': 'elect', 'a': 'A'},
    {'q': 'The sales figures for this quarter are ___ higher than expected.', 'A': 'consider', 'B': 'considerable', 'C': 'considerably', 'D': 'consideration', 'a': 'C'},
    {'q': 'All passengers must ___ their boarding passes at the gate.', 'A': 'present', 'B': 'presented', 'C': 'presenting', 'D': 'presentation', 'a': 'A'},
    {'q': 'The supervisor ___ the team about the schedule change yesterday.', 'A': 'inform', 'B': 'informs', 'C': 'informed', 'D': 'informing', 'a': 'C'},
    {'q': 'The package should arrive ___ three to five business days.', 'A': 'within', 'B': 'during', 'C': 'while', 'D': 'between', 'a': 'A'},
    {'q': 'Ms. Chen has been ___ as the new regional manager.', 'A': 'appoint', 'B': 'appointed', 'C': 'appointing', 'D': 'appointment', 'a': 'B'},
    {'q': 'The client expressed ___ with the quality of service.', 'A': 'satisfy', 'B': 'satisfying', 'C': 'satisfied', 'D': 'satisfaction', 'a': 'D'},
    {'q': 'Please ___ the attached document for your reference.', 'A': 'review', 'B': 'reviewed', 'C': 'reviewing', 'D': 'reviews', 'a': 'A'},
    {'q': 'The training program is ___ to all new employees.', 'A': 'avail', 'B': 'available', 'C': 'availability', 'D': 'availing', 'a': 'B'},
    {'q': 'The invoice must be paid ___ 30 days of receipt.', 'A': 'within', 'B': 'until', 'C': 'before', 'D': 'since', 'a': 'A'},
    {'q': 'We are ___ seeking candidates for the marketing position.', 'A': 'active', 'B': 'actively', 'C': 'activate', 'D': 'activity', 'a': 'B'},
    {'q': 'The report ___ that sales increased by 15 percent.', 'A': 'indicate', 'B': 'indicates', 'C': 'indicating', 'D': 'indication', 'a': 'B'},
    {'q': 'Guests are ___ requested to check out before noon.', 'A': 'kind', 'B': 'kindly', 'C': 'kindness', 'D': 'kinder', 'a': 'B'},
]

for i in range(0, len(toeic_templates), 10):
    batch = toeic_templates[i:i+10]
    qs = [{'index': j, 'question_text': q['q'], 'options': [{'key':'A','text':q['A']},{'key':'B','text':q['B']},{'key':'C','text':q['C']},{'key':'D','text':q['D']}], 'answers': [q['a']], 'explanation': ''} for j, q in enumerate(batch)]
    title = f'TOEIC Part 5 Business {i//10+4}'
    content = {'type': 'grouped_reading', 'group_name': title, 'passage': '', 'direction': 'Choose the best answer to complete each sentence.', 'questions': qs, 'question_count': len(qs)}
    if insert_q('TOEIC', 'reading', 'incomplete_sentences', title, content):
        inserted += 1
        print(f'  {title}: {len(qs)} Qs')

print(f'\n=== TOTAL ENRICHMENT: {inserted} groups ===')
