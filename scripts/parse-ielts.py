"""
IELTS Content Parser — Extracts questions from Cambridge Practice Tests PDF + DOCX worksheets
Inserts into D1 database via Cloudflare API
"""
import sys, json, re, os, urllib.request
sys.stdout.reconfigure(encoding='utf-8')
import fitz  # PyMuPDF
from docx import Document

DB_ID = 'd501b671-128e-4a45-9d90-74b22e6691ce'
ACCOUNT_ID = '6d55097e8961e881b77a4108e1ae40c2'
API_TOKEN = 'cfut_28l0tEkYzywTicfRbbRcRAl8VLuvTDhBfItQPB5Udca04c87'

def query_d1(sql, params=None):
    url = f'https://api.cloudflare.com/client/v4/accounts/{ACCOUNT_ID}/d1/database/{DB_ID}/query'
    body = {'sql': sql}
    if params: body['params'] = params
    req = urllib.request.Request(url, data=json.dumps(body).encode(), headers={
        'Authorization': f'Bearer {API_TOKEN}', 'Content-Type': 'application/json',
    })
    try:
        resp = urllib.request.urlopen(req)
        return json.loads(resp.read())
    except Exception as e:
        print(f'  DB error: {e}')
        return None

def insert_question(test_type, section, question_type, title, content, media_url='', difficulty=3):
    r = query_d1(
        "INSERT INTO test_contents (test_type, section, question_type, title, content, media_url, difficulty, topic, source, status) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
        [test_type, section, question_type, title, json.dumps(content, ensure_ascii=False), media_url, str(difficulty), title, 'curated', 'published']
    )
    return r and r.get('success')

# ============================================================
# STEP 1: Extract from Cambridge Practice Tests PDF
# ============================================================
pdf_path = 'ielts-materials/intermediate/2. Intermediate IELTS/Pre-Test and Post-Test/Cambridge Practice Tests for IELTS 1.pdf'
doc = fitz.open(pdf_path)

full_text = ''
for i in range(doc.page_count):
    text = doc[i].get_text()
    if len(text.strip()) > 20:
        full_text += f'\n--- PAGE {i+1} ---\n{text}'
doc.close()

print(f'Extracted {len(full_text)} chars from Cambridge PDF')

# Split into test sections by finding "Test" headers
# The book has 4 complete tests

# Extract Reading passages with questions
reading_pattern = re.compile(
    r'READING PASSAGE\s*(\d+)(.*?)(?=READING PASSAGE\s*\d+|WRITING|$)',
    re.DOTALL | re.IGNORECASE
)

inserted = 0

for match in reading_pattern.finditer(full_text):
    passage_num = match.group(1)
    passage_content = match.group(2).strip()

    # Split passage from questions
    q_split = re.split(r'Questions?\s+(\d+)\s*[-–]\s*(\d+)', passage_content)

    if len(q_split) < 2:
        continue

    # First part is the passage text
    passage_text = q_split[0].strip()[:3000]  # Limit length

    # Extract question blocks
    questions = []
    for i in range(1, len(q_split) - 2, 3):
        q_start = int(q_split[i])
        q_end = int(q_split[i+1])
        q_block = q_split[i+2].strip()

        # Determine question type from instructions
        q_type = 'multiple_choice'
        if 'TRUE' in q_block and 'FALSE' in q_block and 'NOT GIVEN' in q_block:
            q_type = 'true_false_not_given'
        elif 'YES' in q_block and 'NO' in q_block and 'NOT GIVEN' in q_block:
            q_type = 'true_false_not_given'
        elif 'Complete' in q_block or 'NO MORE THAN' in q_block:
            q_type = 'sentence_completion'
        elif 'Match' in q_block or 'matching' in q_block.lower():
            q_type = 'matching_headings'

        # Extract individual questions (numbered lines)
        q_lines = re.findall(r'(\d+)\s+(.+?)(?=\n\d+\s+|\n\n|$)', q_block)

        for q_num_str, q_text in q_lines:
            q_num = int(q_num_str)
            if q_start <= q_num <= q_end:
                q_text = q_text.strip()[:300]
                if len(q_text) > 5:
                    options = []
                    if q_type == 'true_false_not_given':
                        options = [
                            {'key': 'TRUE', 'text': 'TRUE'},
                            {'key': 'FALSE', 'text': 'FALSE'},
                            {'key': 'NOT GIVEN', 'text': 'NOT GIVEN'},
                        ]
                    questions.append({
                        'index': len(questions),
                        'question_text': q_text,
                        'answers': [],
                        'options': options,
                        'explanation': '',
                    })

    if questions and len(passage_text) > 100:
        content = {
            'type': 'grouped_reading',
            'group_name': f'Reading Passage {passage_num}',
            'passage': passage_text,
            'direction': 'Read the passage and answer the questions.',
            'questions': questions,
            'question_count': len(questions),
        }

        if insert_question('IELTS', 'reading', 'true_false_not_given' if any(q.get('options') for q in questions) else 'sentence_completion',
                          f'Cambridge Test - Passage {passage_num}', content, '', 3):
            inserted += 1
            print(f'  Reading Passage {passage_num}: {len(questions)} questions')

# Extract Listening sections
listening_pattern = re.compile(
    r'SECTION\s*(\d+)\s+Questions?\s+(\d+)\s*[-–]\s*(\d+)(.*?)(?=SECTION\s*\d+|READING|$)',
    re.DOTALL | re.IGNORECASE
)

for match in listening_pattern.finditer(full_text):
    sec_num = match.group(1)
    q_start = int(match.group(2))
    q_end = int(match.group(3))
    sec_content = match.group(4).strip()

    # Determine question type
    q_type = 'fill_in_blank'
    if 'circle' in sec_content.lower() or 'choose' in sec_content.lower():
        q_type = 'multiple_choice'

    # Extract questions
    q_lines = re.findall(r'(\d+)\s*[.)\s]+(.+?)(?=\n\d+\s*[.)\s]+|\n\n|$)', sec_content)
    questions = []
    for q_num_str, q_text in q_lines:
        q_num = int(q_num_str)
        if q_start <= q_num <= q_end:
            q_text = q_text.strip()[:300]
            if len(q_text) > 3:
                questions.append({
                    'index': len(questions),
                    'question_text': q_text,
                    'script': '',
                    'answers': [],
                    'options': [],
                    'explanation': '',
                })

    if questions:
        content = {
            'type': 'grouped_listening',
            'group_name': f'Listening Section {sec_num}',
            'direction': f'Listen to Section {sec_num} and answer questions {q_start}-{q_end}.',
            'passage_script': '',
            'questions': questions,
            'question_count': len(questions),
        }

        if insert_question('IELTS', 'listening', q_type,
                          f'Cambridge Test - Listening Section {sec_num}', content, '', 3):
            inserted += 1
            print(f'  Listening Section {sec_num}: {len(questions)} questions')

print(f'\nPDF extraction done: {inserted} groups inserted')

# ============================================================
# STEP 2: Extract Speaking from DOCX worksheets
# ============================================================
print('\n=== Extracting Speaking from DOCX ===')

speaking_files = []
for root, dirs, files in os.walk('ielts-materials'):
    for f in files:
        if 'speaking' in f.lower() and f.endswith('.docx'):
            speaking_files.append(os.path.join(root, f))

for sf in speaking_files:
    try:
        doc = Document(sf)
        level = 'basic' if 'basic' in sf.lower() else 'intermediate' if 'intermediate' in sf.lower() else 'advanced'
        fname = os.path.basename(sf)

        # Extract speaking prompts
        questions = []
        current_part = ''

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Detect part labels
            if re.match(r'Part\s+[123]', text, re.IGNORECASE):
                current_part = text[:20]

            # Questions are usually sentences ending with ?
            if '?' in text and len(text) > 10 and len(text) < 300:
                questions.append({
                    'index': len(questions),
                    'script': text,
                    'question_text': '',
                    'part': current_part,
                })

        if questions:
            # Group into parts
            part1 = [q for q in questions if 'part 1' in q.get('part', '').lower() or not q.get('part')][:6]
            part2 = [q for q in questions if 'part 2' in q.get('part', '').lower()][:2]
            part3 = [q for q in questions if 'part 3' in q.get('part', '').lower()][:4]

            # Insert Part 1 questions
            if part1:
                content = {
                    'type': 'grouped_speaking',
                    'group_name': f'{level.title()} Speaking Part 1 ({fname})',
                    'direction': 'Answer the following questions about familiar topics. Speak for about 30 seconds per question.',
                    'questions': [{'index': i, 'script': q['script'], 'question_text': ''} for i, q in enumerate(part1)],
                    'question_count': len(part1),
                }
                if insert_question('IELTS', 'speaking', 'part1', f'{level.title()} Part 1 ({fname})', content):
                    inserted += 1
                    print(f'  Speaking Part 1 from {fname}: {len(part1)} questions')

            if part2:
                content = {
                    'type': 'grouped_speaking',
                    'group_name': f'{level.title()} Speaking Part 2 ({fname})',
                    'direction': 'You have 1 minute to prepare. Then speak for 1-2 minutes.',
                    'questions': [{'index': i, 'script': q['script'], 'question_text': ''} for i, q in enumerate(part2)],
                    'question_count': len(part2),
                }
                if insert_question('IELTS', 'speaking', 'part2', f'{level.title()} Part 2 ({fname})', content):
                    inserted += 1

            if part3:
                content = {
                    'type': 'grouped_speaking',
                    'group_name': f'{level.title()} Speaking Part 3 ({fname})',
                    'direction': 'Discuss the following questions in detail.',
                    'questions': [{'index': i, 'script': q['script'], 'question_text': ''} for i, q in enumerate(part3)],
                    'question_count': len(part3),
                }
                if insert_question('IELTS', 'speaking', 'part3', f'{level.title()} Part 3 ({fname})', content):
                    inserted += 1

    except Exception as e:
        print(f'  Error reading {sf}: {e}')

# ============================================================
# STEP 3: Extract Writing from DOCX worksheets
# ============================================================
print('\n=== Extracting Writing from DOCX ===')

writing_files = []
for root, dirs, files in os.walk('ielts-materials'):
    for f in files:
        if 'writing' in f.lower() and (f.endswith('.docx') or f.endswith('.doc')):
            writing_files.append(os.path.join(root, f))

for wf in writing_files:
    if not wf.endswith('.docx'):
        continue
    try:
        doc = Document(wf)
        level = 'basic' if 'basic' in wf.lower() else 'intermediate' if 'intermediate' in wf.lower() else 'advanced'
        fname = os.path.basename(wf)

        # Extract writing prompts
        prompts = []
        current_text = ''

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                if current_text and len(current_text) > 30:
                    prompts.append(current_text)
                    current_text = ''
                continue
            current_text += ' ' + text

        if current_text and len(current_text) > 30:
            prompts.append(current_text)

        # Find task-like prompts (contain "write", "describe", "discuss", etc.)
        task_prompts = [p.strip() for p in prompts if any(kw in p.lower() for kw in ['write', 'describe', 'discuss', 'summarise', 'essay', 'report', 'letter'])]

        for i, prompt in enumerate(task_prompts[:3]):
            task_type = 'task1' if any(kw in prompt.lower() for kw in ['graph', 'chart', 'diagram', 'table', 'describe', 'summarise', 'report']) else 'task2'

            content = {
                'type': 'grouped_writing',
                'group_name': f'{level.title()} Writing ({fname})',
                'direction': prompt[:500],
                'questions': [{
                    'index': 0,
                    'passage': prompt[:1000],
                    'question_text': f'{"Describe the data. Write at least 150 words." if task_type == "task1" else "Write at least 250 words."}',
                    'model_answer': '',
                    'illustrated_passages': [],
                }],
                'question_count': 1,
            }

            if insert_question('IELTS', 'writing', task_type, f'{level.title()} {task_type} ({fname})', content):
                inserted += 1
                print(f'  Writing {task_type} from {fname}')

    except Exception as e:
        print(f'  Error reading {wf}: {e}')

# ============================================================
# STEP 4: Map audio files
# ============================================================
print('\n=== Mapping Audio Files ===')

audio_files = []
for root, dirs, files in os.walk('ielts-materials'):
    for f in files:
        if f.endswith('.mp3'):
            path = os.path.join(root, f)
            level = 'basic' if 'basic' in path.lower() else 'intermediate' if 'intermediate' in path.lower() else 'advanced'
            audio_files.append({'path': path, 'name': f, 'level': level})

print(f'Found {len(audio_files)} audio files:')
for af in audio_files:
    print(f'  [{af["level"]}] {af["name"]}')

print(f'\n=== TOTAL INSERTED: {inserted} IELTS question groups ===')
